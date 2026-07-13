"""
DOCX text loader using python-docx.
Extracts text from paragraphs and tables.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def load_docx(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in parts:
                    parts.append(cell_text)

    return "\n".join(parts).strip()
