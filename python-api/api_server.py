"""
PrepPilot FastAPI Parse & Analysis Microservice.
Runs on port 8001 (Node.js backend on 3000, React frontend on 5173).

Start with:
    python -m uvicorn api_server:app --port 8001 --reload

Agent pipeline (execution order):
    Wave 1 — Parallel: ResumeParser + JDParser
    Wave 2 — Parallel: SkillMatcher + ExperienceEvaluator + EducationEvaluator + CultureFit + ResumeSummarizer
    Wave 3 — Parallel: DecisionAgent + InterviewQuestionAgent (needs skill gaps)
"""
from __future__ import annotations

import asyncio
import io
import sys
import time
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any, Callable, TypeVar

import fitz  # PyMuPDF
import uvicorn
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware

# ── Bootstrap ──────────────────────────────────────────────────────────────────
load_dotenv(dotenv_path=Path(__file__).parent / ".env", override=True)
sys.path.insert(0, str(Path(__file__).parent))

from src.agents.culture_fit import CultureFitAgent
from src.agents.decision_agent import DecisionAgent
from src.agents.education_evaluator import EducationEvaluatorAgent
from src.agents.experience_evaluator import ExperienceEvaluatorAgent
from src.agents.interview_question_agent import InterviewQuestionAgent
from src.agents.jd_parser import JDParserAgent
from src.agents.resume_parser import ResumeParserAgent
from src.agents.resume_summarizer import ResumeSummarizerAgent
from src.agents.skill_matcher import SkillMatcherAgent

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("preppilot")

# ── App ────────────────────────────────────────────────────────────────────────
app = FastAPI(title="PrepPilot Analysis API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:5174",
    ],
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)

_EXECUTOR = ThreadPoolExecutor(max_workers=8)
T = TypeVar("T")

# ── Retry & Threading Helpers ──────────────────────────────────────────────────

def _run_with_retry(fn: Callable[..., T], *args: Any, max_retries: int = 2) -> T:
    """
    Call fn(*args). If a Gemini RESOURCE_EXHAUSTED (quota) error is raised,
    parse the retryDelay from the error message, sleep, and retry once.
    """
    for attempt in range(max_retries + 1):
        try:
            return fn(*args)
        except Exception as exc:
            exc_str = str(exc)
            is_quota = "RESOURCE_EXHAUSTED" in exc_str or "Quota exceeded" in exc_str
            if is_quota and attempt < max_retries:
                # Try to parse the retry delay from the error body
                match = re.search(r"retryDelay.*?(\d+)s", exc_str)
                wait = int(match.group(1)) + 2 if match else 30
                logger.warning(
                    f"Gemini quota hit on attempt {attempt+1}. "
                    f"Waiting {wait}s before retry…"
                )
                time.sleep(wait)
                continue
            raise
    raise RuntimeError("Unreachable")

def _run_in_thread(fn: Callable[..., Any], *args: Any):
    """Run fn(*args) in the thread pool, wrapped with retry logic."""
    loop = asyncio.get_event_loop()
    # We wrap fn in a lambda/partial so _run_with_retry executes it
    return loop.run_in_executor(_EXECUTOR, _run_with_retry, fn, *args)


# ── Text extraction helpers ────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in parts:
                    parts.append(cell_text)
    return "\n".join(parts).strip()


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    elif ext in (".docx", ".doc"):
        return _extract_docx(data)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload PDF or DOCX.")


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "service": "preppilot-analysis-api", "version": "1.0.0"}


@app.post("/parse")
async def analyse_candidate(
    resume_file: UploadFile = File(..., description="Resume PDF or DOCX"),
    jd_text: str = Form(..., description="Raw job description text"),
):
    """
    Full multi-agent candidate analysis pipeline.

    Returns:
        resume, jd, skills, experience, education, culture,
        decision, interview_questions, narrative_summary, weighted_score
    """

    # ── Validate & extract ──────────────────────────────────────────────────────
    filename = resume_file.filename or "resume"
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(400, detail=f"Unsupported file type '{ext}'. Upload PDF or DOCX.")

    file_bytes = await resume_file.read()
    if len(file_bytes) > 10 * 1024 * 1024:
        raise HTTPException(413, detail="File too large. Maximum size is 10 MB.")

    try:
        resume_text = extract_text(filename, file_bytes)
    except ValueError as e:
        raise HTTPException(400, detail=str(e))
    except Exception as e:
        raise HTTPException(422, detail=f"Failed to extract text: {e}")

    if not resume_text.strip():
        raise HTTPException(422, detail="No text could be extracted from the resume.")
    if not jd_text.strip():
        raise HTTPException(400, detail="Job description text cannot be empty.")

    # ── Wave 1: Parse resume + JD in parallel ──────────────────────────────────
    resume_agent = ResumeParserAgent()
    jd_agent = JDParserAgent()

    parsed_resume, parsed_jd = await asyncio.gather(
        _run_in_thread(resume_agent.run, resume_text),
        _run_in_thread(jd_agent.run, jd_text),
    )

    resume_dict = parsed_resume.model_dump()
    jd_dict = parsed_jd.model_dump()

    # ── Wave 2: All evaluators in parallel ────────────────────────────────────
    skill_agent = SkillMatcherAgent()
    exp_agent = ExperienceEvaluatorAgent()
    edu_agent = EducationEvaluatorAgent()
    culture_agent = CultureFitAgent()
    summarizer = ResumeSummarizerAgent()

    (
        skill_result,
        exp_result,
        edu_result,
        culture_result,
        narrative_summary,
    ) = await asyncio.gather(
        _run_in_thread(
            skill_agent.run,
            parsed_resume.skills,
            parsed_jd.required_skills,
            parsed_jd.preferred_skills,
        ),
        _run_in_thread(exp_agent.run, resume_dict, jd_dict),
        _run_in_thread(edu_agent.run, resume_dict, jd_dict),
        _run_in_thread(culture_agent.run, parsed_resume.summary, jd_text),
        _run_in_thread(summarizer.run, resume_dict),
    )

    # ── Compute weighted score ─────────────────────────────────────────────────
    weighted_score = DecisionAgent.compute_weighted_score(
        skill_score=skill_result.score,
        experience_score=exp_result.score,
        education_score=edu_result.score,
        culture_score=culture_result.overall_score,
    )

    # ── Wave 3: Decision + Interview Questions in parallel ────────────────────
    decision_agent = DecisionAgent()
    iq_agent = InterviewQuestionAgent()

    decision_result, iq_result = await asyncio.gather(
        _run_in_thread(
            decision_agent.run,
            resume_dict,
            jd_dict,
            skill_result.matched,
            skill_result.missing,
            skill_result.partial,
            skill_result.score,
            skill_result.reasoning,
            exp_result.score,
            exp_result.reasoning,
            edu_result.score,
            edu_result.reasoning,
            culture_result.overall_score,
            culture_result.reasoning,
            weighted_score,
        ),
        _run_in_thread(
            iq_agent.run,
            skill_result.missing,
            skill_result.partial,
            resume_dict,
            jd_dict,
        ),
    )

    # ── Build interview_questions in the frontend's expected format ─────────────
    # The frontend expects { technical, behavioral, cultural } arrays.
    # We bucket the 10 questions by their position (4 technical, 2 behavioral,
    # 2 verification → technical, 1 scenario → technical, 1 growth → cultural).
    all_questions = iq_result.questions
    technical_qs = all_questions[0:6]   # questions 1-6 (4 tech + 2 verification + 1 scenario)
    behavioral_qs = all_questions[6:8]  # questions 7-8
    cultural_qs = all_questions[8:10]   # questions 9-10

    # ── Return full report ─────────────────────────────────────────────────────
    return {
        # Raw parsed data
        "resume": resume_dict,
        "jd": jd_dict,
        # Narrative
        "narrative_summary": narrative_summary,
        # Individual scores
        "skills": skill_result.model_dump(),
        "experience": {
            "score": exp_result.score,
            "reasoning": exp_result.reasoning,
            "strengths": [],   # Experience evaluator returns score+reasoning; strengths/weaknesses come from decision
            "weaknesses": [],
            "seniority_fit": "well-matched",
        },
        "education": {
            "score": edu_result.score,
            "reasoning": edu_result.reasoning,
            "meets_requirement": edu_result.score >= 50,
        },
        "culture": culture_result.model_dump(),
        # Decision
        "weighted_score": weighted_score,
        "decision": {
            **decision_result.model_dump(),
            "confidence": int(decision_result.confidence * 100),  # convert 0-1 → 0-100
        },
        # Interview questions split into categories for frontend tabs
        "interview_questions": {
            "technical": technical_qs,
            "behavioral": behavioral_qs,
            "cultural": cultural_qs,
        },
    }


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    uvicorn.run("api_server:app", host="0.0.0.0", port=8001, reload=True)
